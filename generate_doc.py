"""Script pour générer le document Word explicatif du projet MOT Pipeline."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.opc.constants


def set_heading_color(heading, r, g, b):
    for run in heading.runs:
        run.font.color.rgb = RGBColor(r, g, b)


def add_formula_box(doc, formula_text, explanation=None):
    """Ajoute un bloc formule en gris avec bordure."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F0F0F0')
    p._p.get_or_add_pPr().append(shading)
    run = p.add_run(formula_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run.bold = True
    if explanation:
        exp = doc.add_paragraph()
        exp.paragraph_format.left_indent = Cm(1.2)
        exp.paragraph_format.space_before = Pt(2)
        r = exp.add_run(f"→ {explanation}")
        r.font.color.rgb = RGBColor(80, 80, 80)
        r.font.size = Pt(10)
        r.italic = True


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'FFF3CD')
    p._p.get_or_add_pPr().append(shading)
    icon = p.add_run("💡 Note : ")
    icon.bold = True
    icon.font.size = Pt(10)
    content = p.add_run(text)
    content.font.size = Pt(10)


def build_document():
    doc = Document()

    # ── Styles globaux ──────────────────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    for h_name, size, bold in [
        ('Heading 1', 18, True),
        ('Heading 2', 14, True),
        ('Heading 3', 12, True),
    ]:
        s = doc.styles[h_name]
        s.font.size = Pt(size)
        s.font.bold = bold

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # ═══════════════════════════════════════════════════════════════════════════
    # PAGE DE TITRE
    # ═══════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Pipeline de Pistage Multi-Objets (MOT)")
    tr.bold = True
    tr.font.size = Pt(24)
    tr.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("Documentation technique — Introduction pour débutants")
    sr.italic = True
    sr.font.size = Pt(14)
    sr.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("Projet ENS — Avril 2026").font.size = Pt(11)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # TABLE DES MATIÈRES (manuelle)
    # ═══════════════════════════════════════════════════════════════════════════
    h = doc.add_heading("Table des matières", level=1)
    set_heading_color(h, 0x1F, 0x49, 0x7D)

    toc_items = [
        ("1.", "Qu'est-ce que le Multi-Object Tracking (MOT) ?"),
        ("2.", "Vue d'ensemble du pipeline"),
        ("3.", "Détection d'objets — YOLOv8"),
        ("4.", "Filtre de Kalman — prédire la position"),
        ("5.", "ByteTrack — l'algorithme de pistage"),
        ("6.", "Re-identification (Re-ID) — reconnaître les individus"),
        ("7.", "Compteur de ligne virtuelle"),
        ("8.", "Évaluation — métriques MOT17"),
        ("9.", "Configuration du pipeline"),
        ("10.", "Structure du projet et extensibilité"),
    ]
    for num, title_text in toc_items:
        p = doc.add_paragraph(style='List Number')
        p.clear()
        run = p.add_run(f"{num}  {title_text}")
        run.font.size = Pt(11)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # 1. QU'EST-CE QUE LE MOT ?
    # ═══════════════════════════════════════════════════════════════════════════
    h = doc.add_heading("1. Qu'est-ce que le Multi-Object Tracking (MOT) ?", level=1)
    set_heading_color(h, 0x1F, 0x49, 0x7D)

    doc.add_paragraph(
        "Le Multi-Object Tracking (MOT) est un domaine de la vision par ordinateur qui répond à "
        "une question simple : étant donné une vidéo, comment localiser et suivre plusieurs "
        "objets — personnes, voitures, animaux — tout au long de la séquence, "
        "même quand ils se croisent, disparaissent temporairement ou sont partiellement cachés ?"
    )

    h2 = doc.add_heading("1.1 Pourquoi est-ce difficile ?", level=2)
    set_heading_color(h2, 0x2E, 0x74, 0xB5)

    challenges = [
        ("Occlusions", "Un piéton peut être caché par une voiture pendant plusieurs secondes."),
        ("Objets similaires", "Dans une foule, deux personnes habillées pareil sont difficiles à distinguer."),
        ("Changements rapides", "À 30 images/seconde, un objet peut se déplacer de nombreux pixels entre deux frames."),
        ("Apparitions/disparitions", "Des personnes entrent et sortent du champ de la caméra."),
        ("Faux positifs / faux négatifs", "Le détecteur peut manquer un objet ou en détecter là où il n'y en a pas."),
    ]
    for name, desc in challenges:
        p = doc.add_paragraph(style='List Bullet')
        p.clear()
        run_bold = p.add_run(f"{name} : ")
        run_bold.bold = True
        p.add_run(desc)

    h2 = doc.add_heading("1.2 Applications concrètes", level=2)
    set_heading_color(h2, 0x2E, 0x74, 0xB5)

    apps = [
        "Surveillance vidéo (comptage de personnes, détection d'anomalies)",
        "Conduite autonome (piétons, cyclistes, véhicules)",
        "Analyse sportive (trajectoires des joueurs)",
        "Robotique (navigation dans des espaces peuplés)",
        "Analyse du trafic routier",
    ]
    for a in apps:
        doc.add_paragraph(a, style='List Bullet')

    h2 = doc.add_heading("1.3 Approche générale : Tracking-by-Detection", level=2)
    set_heading_color(h2, 0x2E, 0x74, 0xB5)

    doc.add_paragraph(
        "Ce projet utilise le paradigme «\u00a0Tracking-by-Detection\u00a0» : on détecte les objets "
        "indépendamment dans chaque image (frame), puis on relie ces détections entre les frames "
        "pour former des trajectoires cohérentes. Ce paradigme est dominant car il découple "
        "le problème en deux sous-problèmes plus simples."
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # 2. VUE D'ENSEMBLE
    # ═══════════════════════════════════════════════════════════════════════════
    h = doc.add_heading("2. Vue d'ensemble du pipeline", level=1)
    set_heading_color(h, 0x1F, 0x49, 0x7D)

    doc.add_paragraph(
        "Le pipeline traite chaque frame vidéo en quatre étapes successives :"
    )

    steps = [
        ("1 — Détection", "YOLOv8 analyse la frame et retourne une liste de boîtes englobantes "
         "(bounding boxes) avec un score de confiance pour chaque objet détecté."),
        ("2 — Extraction de features Re-ID", "Pour chaque détection à haute confiance, "
         "MobileNetV2 extrait un vecteur de 512 nombres qui décrit l'apparence visuelle de l'objet "
         "(couleur, texture, silhouette). Ce vecteur sert d'empreinte visuelle."),
        ("3 — Pistage (ByteTrack)", "L'algorithme ByteTrack associe les nouvelles détections aux "
         "trajectoires existantes grâce à la similarité geometrique (IoU) et visuelle (Re-ID). "
         "Un filtre de Kalman prédit la position attendue de chaque objet."),
        ("4 — Visualisation", "Les résultats sont affichés sur la vidéo : boîtes annotées avec "
         "l'identifiant de suivi, traînées de mouvement, compteur FPS, et lignes virtuelles."),
    ]

    for title_step, desc_step in steps:
        p = doc.add_paragraph()
        r_bold = p.add_run(f"Étape {title_step} : ")
        r_bold.bold = True
        r_bold.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        p.add_run(desc_step)

    doc.add_paragraph()
    add_formula_box(doc,
        "Frame vidéo  →  [YOLOv8]  →  Détections\n"
        "            →  [MobileNetV2] → Embeddings\n"
        "            →  [ByteTrack + Kalman]  →  Trajectoires\n"
        "            →  [Visualizer]  →  Vidéo annotée",
        "Flux de données simplifié du pipeline"
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # 3. YOLOv8
    # ═══════════════════════════════════════════════════════════════════════════
    h = doc.add_heading("3. Détection d'objets — YOLOv8", level=1)
    set_heading_color(h, 0x1F, 0x49, 0x7D)

    h2 = doc.add_heading("3.1 Qu'est-ce que YOLO ?", level=2)
    set_heading_color(h2, 0x2E, 0x74, 0xB5)

    doc.add_paragraph(
        "YOLO (You Only Look Once) est une famille de réseaux de neurones convolutifs conçus pour "
        "la détection d'objets en temps réel. Le principe fondamental : au lieu de regarder l'image "
        "plusieurs fois avec une fenêtre glissante, YOLO analyse toute l'image en une seule passe "
        "(d'où son nom). Cela le rend extrêmement rapide."
    )
    doc.add_paragraph(
        "YOLOv8, développé par Ultralytics en 2023, est la 8ème génération. Le projet utilise "
        "yolov8n.pt (variante «\u00a0nano\u00a0», la plus légère) pour maximiser la vitesse, "
        "au prix d'une précision légèrement réduite par rapport aux versions plus lourdes "
        "(yolov8s, yolov8m, yolov8l, yolov8x)."
    )

    h2 = doc.add_heading("3.2 Comment fonctionne YOLOv8 ?", level=2)
    set_heading_color(h2, 0x2E, 0x74, 0xB5)

    doc.add_paragraph(
        "YOLOv8 divise l'image en une grille. Pour chaque cellule, il prédit :"
    )
    yolo_preds = [
        "Des boîtes englobantes (bounding boxes) : [x1, y1, x2, y2] en pixels",
        "Un score de confiance : probabilité qu'un objet soit présent (0 à 1)",
        "Des probabilités de classe : probabilité que l'objet soit une personne, une voiture, etc.",
    ]
    for yp in yolo_preds:
        doc.add_paragraph(yp, style='List Bullet')

    doc.add_paragraph(
        "Le modèle est entraîné sur le dataset COCO (Common Objects in Context) qui contient "
        "80 classes d'objets. Dans ce pipeline, seule la classe 0 (personne) est utilisée."
    )

    h2 = doc.add_heading("3.3 Seuils et filtrage", level=2)
    set_heading_color(h2, 0x2E, 0x74, 0xB5)

    doc.add_paragraph(
        "Après la détection brute, deux filtres sont appliqués :"
    )

    doc.add_paragraph(
        "Seuil de confiance (conf_thresh = 0.25) : toute détection avec un score < 0.25 est "
        "ignorée. Baisser ce seuil détecte plus d'objets mais génère plus de faux positifs.",
        style='List Bullet'
    )

    doc.add_paragraph(
        "NMS — Non-Maximum Suppression (iou_thresh = 0.45) : quand plusieurs boîtes se "
        "chevauchent fortement pour le même objet, seule la meilleure est conservée. "
        "Le seuil IoU contrôle à partir de quel chevauchement deux boîtes sont considérées "
        "comme détectant le même objet.",
        style='List Bullet'
    )

    add_formula_box(doc,
        "IoU(A, B) = Aire(A ∩ B) / Aire(A ∪ B)",
        "IoU = Intersection sur Union. Vaut 1 si les boîtes sont identiques, 0 si elles ne se touchent pas."
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # 4. FILTRE DE KALMAN
    # ═══════════════════════════════════════════════════════════════════════════
    h = doc.add_heading("4. Filtre de Kalman — Prédire la position des objets", level=1)
    set_heading_color(h, 0x1F, 0x49, 0x7D)

    h2 = doc.add_heading("4.1 Intuition", level=2)
    set_heading_color(h2, 0x2E, 0x74, 0xB5)

    doc.add_paragraph(
        "Imaginez que vous regardez une voiture dans un rétroviseur. Même si la voiture disparaît "
        "une seconde derrière un camion, vous savez approximativement où elle sera quand elle "
        "réapparaîtra, car vous connaissez sa vitesse et sa direction. C'est exactement ce que "
        "fait le filtre de Kalman : il prédit la position future d'un objet en modélisant "
        "son mouvement."
    )
    doc.add_paragraph(
        "Le filtre de Kalman est un algorithme mathématique qui combine deux sources "
        "d'information :"
    )
    doc.add_paragraph("Le modèle de mouvement : «\u00a0où l'objet devrait-il être d'après sa trajectoire passée\u00a0?»", style='List Bullet')
    doc.add_paragraph("La mesure actuelle : «\u00a0où le détecteur dit-il que l'objet se trouve\u00a0?»", style='List Bullet')

    h2 = doc.add_heading("4.2 Le vecteur d'état (8 dimensions)", level=2)
    set_heading_color(h2, 0x2E, 0x74, 0xB5)

    doc.add_paragraph(
        "Chaque objet suivi est décrit par un vecteur d'état à 8 composantes :"
    )

    add_formula_box(doc,
        "x = [cx, cy, ar, h, vx, vy, var, vh]",
        "Le vecteur d'état complet de l'objet"
    )

    state_desc = [
        ("cx, cy", "Coordonnées du centre de la boîte englobante (en pixels)"),
        ("ar", "Rapport d'aspect : largeur / hauteur de la boîte"),
        ("h", "Hauteur de la boîte (en pixels)"),
        ("vx, vy", "Vitesse du centre en x et y (pixels par frame)"),
        ("var", "Vitesse de changement du rapport d'aspect"),
        ("vh", "Vitesse de changement de la hauteur"),
    ]

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = "Variable"
    hdr[1].text = "Description"
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for var, desc in state_desc:
        row = table.add_row().cells
        row[0].text = var
        row[1].text = desc

    doc.add_paragraph()

    h2 = doc.add_heading("4.3 Étape de prédiction", level=2)
    set_heading_color(h2, 0x2E, 0x74, 0xB5)

    doc.add_paragraph(
        "À chaque nouvelle frame, avant de recevoir les nouvelles détections, "
        "le filtre prédit la nouvelle position de l'objet selon le modèle de "
        "mouvement à vitesse constante :"
    )

    add_formula_box(doc,
        "x̂[t] = F · x[t-1]",
        "F est la matrice de transition. Elle applique : nouvelle_position = ancienne_position + vitesse × dt"
    )

    add_formula_box(doc,
        "P̂[t] = F · P[t-1] · Fᵀ + Q",
        "P est la matrice de covariance (incertitude sur l'état). Q est le bruit de processus."
    )

    doc.add_paragraph(
        "Plus l'objet est grand (grande hauteur h), plus l'incertitude sur sa position "
        "est grande — le modèle tient compte de cela en ajustant les bruits proportionnellement à h."
    )

    h2 = doc.add_heading("4.4 Étape de mise à jour (correction)", level=2)
    set_heading_color(h2, 0x2E, 0x74, 0xB5)

    doc.add_paragraph(
        "Quand une détection est associée à la piste, le filtre corrige sa prédiction :"
    )

    add_formula_box(doc,
        "e = z - H · x̂              (innovation = erreur de prédiction)",
        "z est la mesure [cx, cy, ar, h], H projette l'état sur l'espace de mesure"
    )

    add_formula_box(doc,
        "K = P̂ · Hᵀ · (H · P̂ · Hᵀ + R)⁻¹    (gain de Kalman)",
        "K détermine combien on fait confiance à la mesure vs la prédiction"
    )

    add_formula_box(doc,
        "x[t] = x̂ + K · e           (état mis à jour)",
        "Correction de la prédiction par l'innovation pondérée par le gain"
    )

    add_formula_box(doc,
        "P[t] = P̂ - K · H · P̂      (covariance mise à jour)",
        "L'incertitude diminue après avoir reçu une mesure"
    )

    add_note(doc,
        "Si le gain K est proche de 1, on fait très confiance à la mesure (capteur précis). "
        "Si K ≈ 0, on fait confiance au modèle (capteur bruité). Le filtre optimise "
        "automatiquement cet équilibre."
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # 5. BYTETRACK
    # ═══════════════════════════════════════════════════════════════════════════
    h = doc.add_heading("5. ByteTrack — L'algorithme de pistage", level=1)
    set_heading_color(h, 0x1F, 0x49, 0x7D)

    h2 = doc.add_heading("5.1 Idée générale", level=2)
    set_heading_color(h2, 0x2E, 0x74, 0xB5)

    doc.add_paragraph(
        "ByteTrack (Zhang et al., 2022) est un algorithme de pistage multi-objets qui améliore "
        "les approches classiques en exploitant même les détections à faible confiance. "
        "L'intuition : une détection peu confiante peut quand même être utile pour ne pas perdre "
        "un objet momentanément occulté."
    )

    doc.add_paragraph(
        "L'approche classique jetait toutes les détections sous un seuil de confiance. "
        "ByteTrack les conserve pour une deuxième passe d'association, "
        "ce qui réduit significativement les pertes de piste."
    )

    h2 = doc.add_heading("5.2 Les pistes et leurs états", level=2)
    set_heading_color(h2, 0x2E, 0x74, 0xB5)

    doc.add_paragraph(
        "Chaque objet suivi est représenté par une «\u00a0piste\u00a0» (track) qui peut être dans l'un des états suivants :"
    )

    states = [
        ("New (Nouveau)", "Vient d'être créé. Pas encore confirmé — on attend min_hits=2 détections consécutives."),
        ("Active (Actif)", "Confirmé, suivi en cours. Affiché à l'écran avec son identifiant."),
        ("Lost (Perdu)", "Plus de détection depuis quelques frames. Maintenu vivant grâce au filtre de Kalman."),
        ("Removed (Supprimé)", "Perdu depuis trop longtemps (> max_lost_age=30 frames). Définitivement éliminé."),
    ]

    for state_name, state_desc in states:
        p = doc.add_paragraph(style='List Bullet')
        p.clear()
        r = p.add_run(f"{state_name} : ")
        r.bold = True
        p.add_run(state_desc)

    h2 = doc.add_heading("5.3 L'association en deux passes", level=2)
    set_heading_color(h2, 0x2E, 0x74, 0xB5)

    doc.add_paragraph("Voici le déroulement complet à chaque frame :")

    steps_bt = [
        ("Prédiction Kalman",
         "Pour chaque piste active ou perdue, le filtre de Kalman prédit la position attendue "
         "à cette frame (avant de recevoir les nouvelles détections)."),
        ("Séparation des détections",
         "Les détections sont divisées en deux groupes :\n"
         "  • Haute confiance : conf ≥ 0.6 (conf_high)\n"
         "  • Faible confiance : 0.1 ≤ conf < 0.6 (conf_low)"),
        ("1ère association — haute confiance",
         "On associe les pistes actives aux détections haute confiance via la matrice de coût "
         "(IoU + Re-ID). L'algorithme hongrois trouve l'association optimale."),
        ("2ème association — faible confiance",
         "Les pistes non associées (actives ou perdues) sont associées aux détections faible confiance "
         "via IoU uniquement (pas de Re-ID pour les objets peu confiants)."),
        ("Gestion des pistes",
         "• Piste non associée → état Lost (temps_depuis_update++)\n"
         "• Détection non associée → nouvelle piste (état New)\n"
         "• Piste Lost trop longtemps → état Removed"),
    ]

    for i, (step_name, step_desc) in enumerate(steps_bt, 1):
        p = doc.add_paragraph()
        r = p.add_run(f"Étape {i} — {step_name} : ")
        r.bold = True
        r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        p.add_run(step_desc)

    h2 = doc.add_heading("5.4 La matrice de coût", level=2)
    set_heading_color(h2, 0x2E, 0x74, 0xB5)

    doc.add_paragraph(
        "Pour associer pistes et détections, on calcule un «\u00a0coût\u00a0» pour chaque paire "
        "(piste i, détection j). Ce coût combine deux métriques :"
    )

    add_formula_box(doc,
        "coût(i, j)  =  α · (1 − IoU)  +  (1 − α) · distance_cosinus",
        "α = reid_alpha = 0.8 par défaut (80% poids sur la géométrie, 20% sur l'apparence)"
    )

    add_formula_box(doc,
        "IoU(piste_i, détection_j) = Aire(intersection) / Aire(union)",
        "Mesure le chevauchement géométrique. Vaut 1 si les boîtes coïncident."
    )

    add_formula_box(doc,
        "distance_cosinus(emb_i, emb_j)  =  1 − (emb_i · emb_j)",
        "Pour des embeddings normalisés (‖emb‖=1). Vaut 0 si identiques, 1 si orthogonaux."
    )

    h2 = doc.add_heading("5.5 L'algorithme hongrois", level=2)
    set_heading_color(h2, 0x2E, 0x74, 0xB5)

    doc.add_paragraph(
        "Une fois la matrice de coût calculée, on cherche l'association optimale : "
        "comment assigner chaque piste à au plus une détection (et vice versa) pour "
        "minimiser le coût total ? C'est un problème d'affectation bipartite résolu "
        "par l'algorithme hongrois (algorithme de Kuhn-Munkres, complexité O(n³))."
    )

    add_note(doc,
        "Une association n'est retenue que si son coût est inférieur au seuil. "
        "Si le coût est trop élevé (objets trop éloignés ou trop différents), "
        "la piste et la détection restent non associées."
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # 6. RE-IDENTIFICATION
    # ═══════════════════════════════════════════════════════════════════════════
    h = doc.add_heading("6. Re-identification (Re-ID) — Reconnaître les individus", level=1)
    set_heading_color(h, 0x1F, 0x49, 0x7D)

    h2 = doc.add_heading("6.1 Qu'est-ce que la Re-ID ?", level=2)
    set_heading_color(h2, 0x2E, 0x74, 0xB5)

    doc.add_paragraph(
        "La Re-identification (Re-ID) consiste à reconnaître qu'un objet vu à un instant t₁ "
        "est le même que celui vu à un instant t₂, même si entre-temps il a disparu du champ "
        "de la caméra. On extrait une «\u00a0empreinte visuelle\u00a0» de chaque objet et on compare "
        "ces empreintes pour établir des correspondances."
    )

    h2 = doc.add_heading("6.2 MobileNetV2 comme extracteur de features", level=2)
    set_heading_color(h2, 0x2E, 0x74, 0xB5)

    doc.add_paragraph(
        "Ce projet utilise MobileNetV2, un réseau de neurones léger et efficace conçu pour "
        "fonctionner sur des appareils à ressources limitées. MobileNetV2 a été pré-entraîné "
        "sur ImageNet (1.2 million d'images, 1000 classes) et sert ici de feature extractor."
    )

    doc.add_paragraph("Le traitement d'une détection se déroule ainsi :")

    reid_steps = [
        "Recadrage (crop) de la boîte englobante depuis la frame",
        "Redimensionnement à 224×224 pixels (taille standard d'entrée de MobileNetV2)",
        "Normalisation des valeurs de pixels selon les statistiques ImageNet",
        "Passage dans MobileNetV2 → vecteur de 1280 features",
        "Pooling adaptatif global → réduction à un vecteur de dimension variable",
        "Projection linéaire → vecteur de dimension 512 (configurable)",
        "Normalisation L2 → le vecteur est ramené sur la sphère unité ‖emb‖ = 1",
    ]

    for i, step in enumerate(reid_steps, 1):
        doc.add_paragraph(f"{i}. {step}", style='List Number')

    h2 = doc.add_heading("6.3 Pourquoi la normalisation L2 ?", level=2)
    set_heading_color(h2, 0x2E, 0x74, 0xB5)

    doc.add_paragraph(
        "En normalisant les vecteurs à longueur 1, le produit scalaire entre deux vecteurs "
        "devient directement la similarité cosinus. Cela simplifie le calcul et rend la "
        "comparaison indépendante de la magnitude des vecteurs."
    )

    add_formula_box(doc,
        "emb_normalisé  =  emb / ‖emb‖₂",
        "Division de chaque composante par la norme euclidienne du vecteur"
    )

    add_formula_box(doc,
        "similitude(emb_a, emb_b)  =  emb_a · emb_b    (pour des vecteurs normalisés)",
        "Vaut 1 si identiques, 0 si perpendiculaires, -1 si opposés"
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # 7. COMPTEUR DE LIGNE VIRTUELLE
    # ═══════════════════════════════════════════════════════════════════════════
    h = doc.add_heading("7. Compteur de ligne virtuelle", level=1)
    set_heading_color(h, 0x1F, 0x49, 0x7D)

    doc.add_paragraph(
        "Le module VirtualLineCounter permet de compter les objets qui franchissent "
        "une ligne définie dans la configuration (par défaut une ligne horizontale "
        "au centre de l'image). Il distingue deux directions : «\u00a0in\u00a0» et «\u00a0out\u00a0»."
    )

    h2 = doc.add_heading("7.1 Détection de franchissement", level=2)
    set_heading_color(h2, 0x2E, 0x74, 0xB5)

    doc.add_paragraph(
        "On utilise le produit vectoriel pour déterminer de quel côté d'une ligne se trouve "
        "le centre d'un objet. Si le signe change entre deux frames consécutives, "
        "l'objet a franchi la ligne."
    )

    add_formula_box(doc,
        "cross = (P2.x - P1.x) × (C.y - P1.y) - (P2.y - P1.y) × (C.x - P1.x)",
        "P1, P2 : extrémités de la ligne virtuelle  |  C : centre de l'objet"
    )

    add_formula_box(doc,
        "franchissement détecté  ⟺  cross_prev × cross_curr < 0",
        "Le changement de signe indique que l'objet a traversé la ligne"
    )

    doc.add_paragraph(
        "La direction est déterminée par le signe de cross_curr :\n"
        "  • cross_curr < 0  →  direction «\u00a0in\u00a0»\n"
        "  • cross_curr > 0  →  direction «\u00a0out\u00a0»"
    )

    doc.add_paragraph(
        "Les résultats sont exportés en CSV avec : horodatage, identifiant de piste, "
        "classe de l'objet, direction."
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # 8. MÉTRIQUES D'ÉVALUATION
    # ═══════════════════════════════════════════════════════════════════════════
    h = doc.add_heading("8. Évaluation — Métriques MOT17", level=1)
    set_heading_color(h, 0x1F, 0x49, 0x7D)

    doc.add_paragraph(
        "Pour évaluer objectivement les performances du pipeline, on utilise le benchmark "
        "MOT17 (MOT Challenge 2017), un dataset standard comportant des séquences vidéo "
        "annotées avec les positions réelles de chaque personne (ground truth)."
    )

    h2 = doc.add_heading("8.1 MOTA — Multi-Object Tracking Accuracy", level=2)
    set_heading_color(h2, 0x2E, 0x74, 0xB5)

    doc.add_paragraph("La métrique principale du MOT Challenge. Elle pénalise trois types d'erreurs :")
    doc.add_paragraph("FN (False Negatives) : objets réels non détectés", style='List Bullet')
    doc.add_paragraph("FP (False Positives) : détections sans objet réel correspondant", style='List Bullet')
    doc.add_paragraph("IDSW (ID Switches) : changements d'identifiant sur un même objet", style='List Bullet')

    add_formula_box(doc,
        "MOTA = 1 - (FN + FP + IDSW) / GT",
        "GT = nombre total d'objets dans la vérité terrain. Plus MOTA est proche de 1 (100%), mieux c'est."
    )

    doc.add_paragraph(
        "MOTA peut être négatif si les erreurs dépassent le nombre d'objets réels. "
        "Les bons systèmes obtiennent des MOTA entre 50% et 80% sur MOT17."
    )

    h2 = doc.add_heading("8.2 IDF1 — ID F1-score", level=2)
    set_heading_color(h2, 0x2E, 0x74, 0xB5)

    doc.add_paragraph(
        "Mesure la capacité du système à maintenir un identifiant cohérent tout au long "
        "d'une trajectoire. Contrairement à MOTA, IDF1 pénalise fortement les erreurs d'identité "
        "même si les objets sont bien détectés."
    )

    add_formula_box(doc,
        "IDF1 = 2 × IDTP / (GT + DT)",
        "IDTP = ID True Positives (détections correctement associées à leur vrai identifiant)\n"
        "GT = total ground truth  |  DT = total détections"
    )

    h2 = doc.add_heading("8.3 ID Switches (IDSW)", level=2)
    set_heading_color(h2, 0x2E, 0x74, 0xB5)

    doc.add_paragraph(
        "Nombre de fois où l'identifiant attribué à un objet réel change. "
        "Un IDSW se produit typiquement lors d'occlusions : la piste est perdue, "
        "et quand l'objet réapparaît, il obtient un nouvel identifiant. "
        "Un bon tracker minimise ce nombre."
    )

    h2 = doc.add_heading("8.4 Résumé des métriques", level=2)
    set_heading_color(h2, 0x2E, 0x74, 0xB5)

    table2 = doc.add_table(rows=1, cols=4)
    table2.style = 'Table Grid'
    headers2 = ["Métrique", "Signification", "Objectif", "Plage typique"]
    for i, hdr_text in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = hdr_text
        for run in cell.paragraphs[0].runs:
            run.bold = True

    metrics_data = [
        ("MOTA", "Précision globale du pistage", "↑ Maximiser", "0–80%"),
        ("IDF1", "Cohérence des identifiants", "↑ Maximiser", "0–75%"),
        ("IDSW", "Nombre de switch d'identité", "↓ Minimiser", "0–∞"),
        ("FPS", "Frames par seconde", "↑ Maximiser", "10–100 FPS"),
    ]

    for row_data in metrics_data:
        row = table2.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # 9. CONFIGURATION
    # ═══════════════════════════════════════════════════════════════════════════
    h = doc.add_heading("9. Configuration du pipeline", level=1)
    set_heading_color(h, 0x1F, 0x49, 0x7D)

    doc.add_paragraph(
        "Tous les paramètres sont centralisés dans config/default.yaml. "
        "Voici les paramètres clés et leur impact :"
    )

    config_table = doc.add_table(rows=1, cols=4)
    config_table.style = 'Table Grid'
    headers3 = ["Paramètre", "Valeur par défaut", "Rôle", "Impact si modifié"]
    for i, hdr_text in enumerate(headers3):
        cell = config_table.rows[0].cells[i]
        cell.text = hdr_text
        for run in cell.paragraphs[0].runs:
            run.bold = True

    config_data = [
        ("conf_thresh", "0.25", "Seuil de confiance détecteur", "↓ = plus de détections, plus de FP"),
        ("iou_thresh (NMS)", "0.45", "NMS du détecteur", "↑ = moins de suppression de doublons"),
        ("conf_high", "0.60", "Seuil haute confiance ByteTrack", "↑ = plus sélectif en 1re passe"),
        ("conf_low", "0.10", "Seuil basse confiance ByteTrack", "↓ = plus d'infos en 2e passe"),
        ("max_lost_age", "30", "Frames avant suppression piste", "↑ = moins de pertes sur occlusions"),
        ("min_hits", "2", "Détections pour confirmer piste", "↑ = moins de faux tracks"),
        ("reid_alpha", "0.80", "Poids IoU vs Re-ID dans le coût", "↓ = plus de poids sur l'apparence"),
        ("dim (embedder)", "512", "Dimension du vecteur Re-ID", "↑ = plus discriminant mais plus lent"),
    ]

    for row_data in config_data:
        row = config_table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # 10. STRUCTURE DU PROJET
    # ═══════════════════════════════════════════════════════════════════════════
    h = doc.add_heading("10. Structure du projet et extensibilité", level=1)
    set_heading_color(h, 0x1F, 0x49, 0x7D)

    h2 = doc.add_heading("10.1 Organisation des fichiers", level=2)
    set_heading_color(h2, 0x2E, 0x74, 0xB5)

    add_formula_box(doc,
        "Pipeline-Multi-Object-Tracking/\n"
        "├── config/\n"
        "│   ├── default.yaml          # Tous les paramètres du pipeline\n"
        "│   └── virtual_line.json     # Coordonnées de la ligne virtuelle\n"
        "├── mot_pipeline/\n"
        "│   ├── pipeline.py           # Orchestrateur principal\n"
        "│   ├── types.py              # Structures de données (Detection, etc.)\n"
        "│   ├── visualizer.py         # Affichage vidéo\n"
        "│   ├── counter.py            # Compteur de ligne virtuelle\n"
        "│   ├── benchmark.py          # Évaluation MOT17\n"
        "│   ├── detector/\n"
        "│   │   └── yolo_detector.py  # YOLOv8\n"
        "│   ├── tracker/\n"
        "│   │   ├── bytetrack.py      # Algorithme ByteTrack\n"
        "│   │   ├── kalman.py         # Filtre de Kalman 8D\n"
        "│   │   └── track.py          # Classe Track + états\n"
        "│   └── reid/\n"
        "│       └── embedder.py       # MobileNetV2 Re-ID\n"
        "└── tests/                    # Tests unitaires"
    )

    h2 = doc.add_heading("10.2 Extensibilité", level=2)
    set_heading_color(h2, 0x2E, 0x74, 0xB5)

    doc.add_paragraph(
        "Le projet utilise des classes abstraites pour permettre de remplacer facilement "
        "chaque composant sans modifier le reste du code :"
    )

    ext_items = [
        ("Remplacer YOLOv8", "Implémenter BaseDetector avec detect(frame) → list[Detection]",
         "RT-DETR, EfficientDet, Faster R-CNN..."),
        ("Remplacer MobileNetV2", "Implémenter BaseEmbedder avec embed(crops) → np.ndarray",
         "OSNet, ResNet-50, ViT..."),
        ("Remplacer ByteTrack", "Implémenter un nouveau tracker dans mot_pipeline/tracker/",
         "DeepSORT, StrongSORT, OC-SORT..."),
    ]

    for ext_name, how, examples in ext_items:
        p = doc.add_paragraph(style='List Bullet')
        p.clear()
        r_name = p.add_run(f"{ext_name} : ")
        r_name.bold = True
        p.add_run(f"{how}. Exemples alternatifs : {examples}.")

    h2 = doc.add_heading("10.3 Commandes principales", level=2)
    set_heading_color(h2, 0x2E, 0x74, 0xB5)

    commands = [
        ("Lancer sur webcam", "python -m mot_pipeline --source 0"),
        ("Lancer sur vidéo", "python -m mot_pipeline --source video.mp4"),
        ("Avec ligne virtuelle", "python -m mot_pipeline --source video.mp4 --line config/virtual_line.json"),
        ("Mode sans affichage", "python -m mot_pipeline --source video.mp4 --no-display"),
        ("Benchmark MOT17", "python run_benchmark.py"),
        ("Tests unitaires", "pytest tests/ -v"),
    ]

    for desc_cmd, cmd in commands:
        p = doc.add_paragraph(style='List Bullet')
        p.clear()
        r_desc = p.add_run(f"{desc_cmd} : ")
        r_desc.bold = True
        r_code = p.add_run(cmd)
        r_code.font.name = 'Courier New'
        r_code.font.size = Pt(10)

    # ─── CONCLUSION ─────────────────────────────────────────────────────────
    doc.add_page_break()
    h = doc.add_heading("Conclusion", level=1)
    set_heading_color(h, 0x1F, 0x49, 0x7D)

    doc.add_paragraph(
        "Ce pipeline MOT illustre l'intégration de plusieurs technologies de vision par "
        "ordinateur dans un système cohérent :"
    )

    conclusions = [
        "YOLOv8 fournit des détections rapides et précises grâce à son architecture neuronale de détection en une passe.",
        "Le filtre de Kalman assure la continuité des trajectoires même lors d'occlusions courtes, en modélisant la dynamique des objets.",
        "ByteTrack exploite intelligemment les détections à faible confiance pour réduire les pertes de piste.",
        "MobileNetV2 ajoute une dimension d'apparence visuelle permettant de réassocier des pistes après de longues occlusions.",
        "L'évaluation sur MOT17 permet de quantifier objectivement les performances et de comparer aux approches de l'état de l'art.",
    ]

    for c in conclusions:
        doc.add_paragraph(c, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph(
        "L'architecture modulaire du projet (classes abstraites, configuration YAML centralisée) "
        "facilite l'expérimentation : chaque composant peut être remplacé indépendamment pour "
        "évaluer de nouvelles approches sans réécrire l'ensemble du système."
    )

    return doc


if __name__ == "__main__":
    doc = build_document()
    output_path = "Documentation_MOT_Pipeline.docx"
    doc.save(output_path)
    print(f"Document généré : {output_path}")
